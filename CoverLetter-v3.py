from docx import Document

first_name = input("What is your first name?: ")
second_name = input("What is your surname?: ")
title = input("What is the Job Title you are applying for: ")
company_name = input("What is the name of your company: ")
previous_role = input("What would your previous role (Job Title): ")
years_experience = input("How many years experience do you have?: ")
projects_completed = input("What projects have you completed?: ")
file_name = input("What do you want to call this file?: ")

sentence = f"To whom it may concern, \n\nI wish to apply for the role of {title} at {company_name} as a highly competent {title} i would bring a dedicated, reliable and driven attitude to this role. \nAs you can see from my attached CV, i have over {years_experience} years experience in the industry and I believe the knowledge and skills i built up during this time make me the perfect candidate for the role. \nI pride myself on my ability to work autonomously, or as part of a larger group. \nIn my previous role as a {previous_role} i would demonstrate being able to lead and/or contribute in a collaborative fashion. I also managed to complete projects such as {projects_completed} \nI hope you find me suitable for the role and i look forward to hearing from you soon. \n\nKind regards"

document = Document()

document.add_heading(f"Cover Letter for  {company_name} role", level=1)

p = document.add_paragraph(sentence)

#p1 = document.add_paragraph((first_name) + " " + (second_name), style="Caption")

document.save(file_name + ".docx")